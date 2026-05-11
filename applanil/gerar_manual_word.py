from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_colored_heading(doc, text, level=1, color=None):
    """Add a heading with optional color"""
    heading = doc.add_heading(text, level=level)
    if color:
        for run in heading.runs:
            run.font.color.rgb = color
    return heading

def shade_paragraph(paragraph, color):
    """Add background color to paragraph"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    paragraph._element.get_or_add_pPr().append(shading_elm)

# Criar documento
doc = Document()

# Configurar margens
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# ============ CAPA ============
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("📘 MANUAL DE USUÁRIO")
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(26, 60, 110)  # Azul escuro

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Tabela de Amortização de Financiamento Imobiliário")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(46, 117, 182)  # Azul médio

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run("Versão 1.0\nAbril 2026")
run.font.size = Pt(11)
run.font.italic = True

doc.add_page_break()

# ============ SUMÁRIO ============
add_colored_heading(doc, "📋 Sumário", level=1, color=RGBColor(26, 60, 110))
toc_items = [
    "Visão Geral",
    "Como Começar",
    "Seção de Configurações",
    "Preenchimento da Tabela",
    "Interpretando os Resultados",
    "Dicas Avançadas",
    "Perguntas Frequentes"
]
for item in toc_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============ VISÃO GERAL ============
add_colored_heading(doc, "🏠 Visão Geral", level=1, color=RGBColor(26, 60, 110))

p = doc.add_paragraph()
run = p.add_run("A ")
run = p.add_run("Tabela de Amortização de Financiamento Imobiliário")
run.font.bold = True
run = p.add_run(" é uma planilha interativa que permite:")

recursos = [
    "Calcular parcelas automaticamente (sistemas PRICE e SAC)",
    "Acompanhar pagamentos mês a mês",
    "Controlar juros e amortizações adicionais",
    "Visualizar saldo devedor em tempo real",
    "Gerar relatórios de quitação e progresso"
]
for recurso in recursos:
    doc.add_paragraph(recurso, style='List Bullet')

add_colored_heading(doc, "Recursos Principais", level=2, color=RGBColor(46, 117, 182))
recursos_principais = [
    "360 meses de previsão (30 anos)",
    "Dois sistemas de cálculo: PRICE (parcela fixa) e SAC (parcela decrescente)",
    "Amortização adicional para acelerar quitação",
    "Seguro mensal configurável",
    "Formatação visual com cores para fácil interpretação",
    "Totalizadores automáticos"
]
for recurso in recursos_principais:
    doc.add_paragraph(recurso, style='List Bullet')

doc.add_page_break()

# ============ COMO COMEÇAR ============
add_colored_heading(doc, "🚀 Como Começar", level=1, color=RGBColor(26, 60, 110))

add_colored_heading(doc, "Passo 1: Abrir a Planilha", level=2)
doc.add_paragraph("Execute o arquivo create_excel_imovel.py", style='List Number')
doc.add_paragraph("Será gerado o arquivo Imovel_Financiado_PERSONALIZADO.xlsx", style='List Number')
doc.add_paragraph("Abra o arquivo no Excel, Google Sheets ou LibreOffice Calc", style='List Number')

add_colored_heading(doc, "Passo 2: Entender a Estrutura", level=2)
doc.add_paragraph("A planilha possui três áreas principais:")
doc.add_paragraph("Zona Roxa (linhas 3-7): Configurações do financiamento", style='List Bullet')
doc.add_paragraph("Zona Verde (linhas 10-11): Totalizadores", style='List Bullet')
doc.add_paragraph("Zona Azul (linhas 13+): Tabela de dados mensal", style='List Bullet')

doc.add_page_break()

# ============ CONFIGURAÇÕES ============
add_colored_heading(doc, "⚙️ Seção de Configurações", level=1, color=RGBColor(26, 60, 110))

p = doc.add_paragraph()
run = p.add_run("As configurações estão na ")
run = p.add_run("seção superior roxa (linhas 4-7)")
run.bold = True

add_colored_heading(doc, "Campos Obrigatórios (lado esquerdo)", level=2)

configs_obrigatorias = [
    ("💰 Valor Financiado (R$)", "500.000,00", "O valor total que você pediu emprestado"),
    ("📅 Mês/Ano de Início", "Jan/2025", "Quando começa seu financiamento"),
    ("⏳ Prazo (meses)", "360", "Quantos meses levará para pagar"),
    ("📈 Taxa de Juros (% a.m.)", "0,75%", "Percentual de juros mensal")
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Campo'
hdr_cells[1].text = 'Exemplo'
hdr_cells[2].text = 'Descrição'

for campo, exemplo, descricao in configs_obrigatorias:
    row_cells = table.add_row().cells
    row_cells[0].text = campo
    row_cells[1].text = exemplo
    row_cells[2].text = descricao

doc.add_paragraph()

add_colored_heading(doc, "Campos Opcionais (lado direito)", level=2)

configs_opcionais = [
    ("🏦 Sistema", "SAC ou PRICE", "Método de cálculo das parcelas"),
    ("🛡️ Seguro Mensal (R$)", "250,00", "Seguro residencial mensal (opcional)"),
    ("📊 Amortização Adicional Padrão", "1.000,00", "Valor extra por padrão em cada mês")
]

table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Light Grid Accent 1'
hdr_cells2 = table2.rows[0].cells
hdr_cells2[0].text = 'Campo'
hdr_cells2[1].text = 'Exemplo'
hdr_cells2[2].text = 'Descrição'

for campo, exemplo, descricao in configs_opcionais:
    row_cells = table2.add_row().cells
    row_cells[0].text = campo
    row_cells[1].text = exemplo
    row_cells[2].text = descricao

add_colored_heading(doc, "Comparação: PRICE vs SAC", level=2)

table3 = doc.add_table(rows=1, cols=3)
table3.style = 'Light Grid Accent 1'
hdr_cells3 = table3.rows[0].cells
hdr_cells3[0].text = 'Aspecto'
hdr_cells3[1].text = 'PRICE'
hdr_cells3[2].text = 'SAC'

comparacoes = [
    ("Parcela", "Fixa", "Diminui"),
    ("Primeira parcela", "Média", "Maior"),
    ("Últimas parcelas", "Média", "Menor"),
    ("Total de juros", "Mais alto", "Mais baixo"),
    ("Melhor para", "Orçamento previsível", "Quer quitar logo")
]

for aspecto, price, sac in comparacoes:
    row_cells = table3.add_row().cells
    row_cells[0].text = aspecto
    row_cells[1].text = price
    row_cells[2].text = sac

doc.add_page_break()

# ============ PREENCHIMENTO DA TABELA ============
add_colored_heading(doc, "📊 Preenchimento da Tabela", level=1, color=RGBColor(26, 60, 110))

colunas_info = [
    ("Coluna A", "Decorativa", "Apenas decoração visual (cores azuis alternadas)"),
    ("Coluna B", "#", "Número da parcela (preenchida automaticamente)"),
    ("Coluna C", "📅 MÊS / ANO", "Preenchida automaticamente"),
    ("Coluna D", "📆 VENCIMENTO", "Data de vencimento (dia 5 de cada mês)"),
    ("Coluna J", "➕ AMORT. ADICIONAL", "⚡ VOCÊ PREENCHE - Valor extra naquele mês"),
    ("Coluna L", "📅 DATA PGTO", "⚡ VOCÊ PREENCHE - Data que você pagou"),
    ("Coluna M", "✏️ VALOR PAGO", "⚡ VOCÊ PREENCHE - Quanto você pagou"),
]

table4 = doc.add_table(rows=1, cols=3)
table4.style = 'Light Grid Accent 1'
hdr_cells4 = table4.rows[0].cells
hdr_cells4[0].text = 'Coluna'
hdr_cells4[1].text = 'Nome'
hdr_cells4[2].text = 'Descrição'

for col, nome, desc in colunas_info:
    row_cells = table4.add_row().cells
    row_cells[0].text = col
    row_cells[1].text = nome
    row_cells[2].text = desc

add_colored_heading(doc, "Campos Para Você Preencher", level=2, color=RGBColor(46, 117, 182))

campos_preenchimento = [
    ("Coluna J - Amortização Adicional", [
        "Deixe em branco: Usa valor padrão de H6",
        "Digite um valor: Usa esse valor naquele mês",
        "Digite 0: Não paga nada de extra"
    ]),
    ("Coluna L - Data de Pagamento", [
        "Deixe em branco: Indica parcela ainda não paga",
        "Digite uma data: Marca quando pagou (DD/MM/YYYY)",
        "Exemplo: 15/01/2025"
    ]),
    ("Coluna M - Valor Pago", [
        "Deixe em branco: Indica não pago",
        "Digite o valor: Quanto você pagou em reais",
        "Pode ser diferente da parcela se pagou com atraso"
    ])
]

for campo, instrucoes in campos_preenchimento:
    add_colored_heading(doc, campo, level=3)
    for instr in instrucoes:
        doc.add_paragraph(instr, style='List Bullet')

doc.add_page_break()

# ============ INTERPRETANDO RESULTADOS ============
add_colored_heading(doc, "📈 Interpretando os Resultados", level=1, color=RGBColor(26, 60, 110))

add_colored_heading(doc, "Zona de Totalizadores (Linha 10-11)", level=2)

totalizadores = [
    ("✅ PARCELAS PAGAS", "Quantas parcelas você já quitou"),
    ("💰 TOTAL AMORTIZADO", "Quanto de principal você já pagou"),
    ("💸 TOTAL JUROS PAGOS", "Quanto você já pagou de juros"),
    ("🛡️ TOTAL SEGUROS", "Total de seguro pago até agora"),
    ("📊 SALDO DEVEDOR ATUAL", "Quanto você ainda deve"),
    ("📈 % QUITADO", "Percentual do financiamento que você quitou")
]

for total, descricao in totalizadores:
    p = doc.add_paragraph()
    run = p.add_run(total)
    run.bold = True
    run = p.add_run(": " + descricao)

add_colored_heading(doc, "Status das Parcelas (Coluna O)", level=2)

status = [
    ("✅ PAGO", "Você pagou o valor total ou mais"),
    ("⚠️ PARCIAL", "Você pagou, mas menos que o devido"),
    ("⏳ PENDENTE", "Ainda não pagou")
]

table5 = doc.add_table(rows=1, cols=2)
table5.style = 'Light Grid Accent 1'
hdr_cells5 = table5.rows[0].cells
hdr_cells5[0].text = 'Status'
hdr_cells5[1].text = 'Significado'

for stat, sig in status:
    row_cells = table5.add_row().cells
    row_cells[0].text = stat
    row_cells[1].text = sig

doc.add_page_break()

# ============ DICAS AVANÇADAS ============
add_colored_heading(doc, "💡 Dicas Avançadas", level=1, color=RGBColor(26, 60, 110))

add_colored_heading(doc, "Como Simular Pagamentos Antecipados", level=2)
doc.add_paragraph("1. Escolha o mês que quer pagar antecipadamente", style='List Number')
doc.add_paragraph("2. Coloque a data na Coluna L", style='List Number')
doc.add_paragraph("3. Coloque o valor na Coluna M", style='List Number')
doc.add_paragraph("4. A situação mudará automaticamente para ✅ PAGO ou ⚠️ PARCIAL", style='List Number')

add_colored_heading(doc, "Como Usar Amortização Adicional Eficientemente", level=2)
doc.add_paragraph("Sistema SAC + Amortização adicional = Melhor combinação", style='List Bullet')
doc.add_paragraph("No sistema SAC, você já paga mais no início", style='List Bullet')
doc.add_paragraph("Adicione amortização extra nos primeiros meses", style='List Bullet')
doc.add_paragraph("Resultado: Juros caem rapidamente, saldo diminui rápido", style='List Bullet')

add_colored_heading(doc, "Dicas para Economizar em Juros", level=2)
doc.add_paragraph("Aumente amortização nos primeiros anos - juros são maiores no início", style='List Bullet')
doc.add_paragraph("Compare PRICE vs SAC - SAC economiza mais juros", style='List Bullet')
doc.add_paragraph("Recebeu bônus ou 13º? Coloque em amortização adicional", style='List Bullet')
doc.add_paragraph("Use a barra de saldo devedor - motivação visual de progresso", style='List Bullet')

doc.add_page_break()

# ============ FAQ ============
add_colored_heading(doc, "❓ Perguntas Frequentes", level=1, color=RGBColor(26, 60, 110))

faqs = [
    ("Qual sistema escolher, PRICE ou SAC?", 
     "PRICE: Se prefere parcela fixa e previsível\nSAC: Se quer pagar menos juros no total"),
    
    ("Posso mudar o sistema depois de preencher dados?",
     "Sim, mas os cálculos de juros mudarão. Todos os dados da tabela se recalcularão automaticamente."),
    
    ("O que fazer se cometi um erro ao preencher?",
     "1. Clique na célula errada\n2. Limpe o valor (Delete)\n3. Digite o valor correto\n4. Tudo se recalcula automaticamente"),
    
    ("Como saber se estou pagando corretamente?",
     "Olhe a Coluna O:\n✅ PAGO = Tudo certo\n⚠️ PARCIAL = Você pagou menos\n⏳ PENDENTE = Ainda não pagou"),
    
    ("Posso usar a planilha para simular diferentes cenários?",
     "Sim! Salve com nome diferente, mude os valores de configuração e compare os cenários lado a lado."),
    
    ("Os valores incluem inflação?",
     "Não. Esta planilha considera taxa de juros fixa, sem inflação.")
]

for pergunta, resposta in faqs:
    p = doc.add_paragraph()
    run = p.add_run("P: " + pergunta)
    run.bold = True
    run.font.color.rgb = RGBColor(131, 56, 236)  # Roxo
    
    p = doc.add_paragraph()
    run = p.add_run("R: " + resposta)
    run.font.color.rgb = RGBColor(45, 45, 45)  # Cinza escuro

doc.add_page_break()

# ============ LEGENDA DE CORES ============
add_colored_heading(doc, "🎨 Legenda de Cores", level=1, color=RGBColor(26, 60, 110))

cores_tabela = [
    ("🟨 Amarelo/Laranja", "Campo para você preencher (entrada de dados)"),
    ("🟢 Verde claro", "Amortização / valores calculados positivos"),
    ("🔴 Vermelho claro", "Juros (diminui com SAC ao longo do tempo)"),
    ("🔵 Azul claro", "Saldo devedor (barra visual)"),
    ("✅ Verde escuro", "Parcela PAGA"),
    ("⏳ Amarelo intenso", "Parcela PENDENTE"),
    ("⚠️ Vermelho intenso", "Parcela PARCIAL")
]

table6 = doc.add_table(rows=1, cols=2)
table6.style = 'Light Grid Accent 1'
hdr_cells6 = table6.rows[0].cells
hdr_cells6[0].text = 'Cor'
hdr_cells6[1].text = 'Significado'

for cor, sig in cores_tabela:
    row_cells = table6.add_row().cells
    row_cells[0].text = cor
    row_cells[1].text = sig

doc.add_page_break()

# ============ CONCLUSÃO ============
add_colored_heading(doc, "✨ Conclusão", level=1, color=RGBColor(26, 60, 110))

conclusao_items = [
    "📊 Acompanhar seu financiamento",
    "💡 Simular cenários",
    "💰 Identificar oportunidades de economia",
    "📈 Visualizar progresso"
]

p = doc.add_paragraph()
run = p.add_run("A Tabela de Amortização é uma ferramenta poderosa para:")

for item in conclusao_items:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_paragraph()
run = p.add_run("Dica final: ")
run.bold = True
run = p.add_run("Atualize a planilha regularmente com seus pagamentos reais. Isso ajuda a manter um controle preciso e identificar se há alguma discrepância com seu banco.")
run.italic = True

p = doc.add_paragraph()
run = p.add_run("Boa sorte com seu financiamento! 🏠")
run.font.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(6, 214, 160)  # Verde

# ============ SALVAR ============
doc.save("Manual_Financiamento_Imobiliario.docx")
print("✅ Manual em Word gerado com sucesso: Manual_Financiamento_Imobiliario.docx")
